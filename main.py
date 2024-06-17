# Libs
import os
import docx
import datetime as dt
import tkinter as tk
from docx2pdf import convert
from tkinter import filedialog, messagebox

class InvoiceMapper:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Invoice Generator")
        self.root.geometry('500x600')

        # Define colors
        label_bg_color = 'lightgrey'    # Background color for labels
        label_fg_color = 'black'        # Foreground (text) color for labels
        entry_bg_color = 'white'        
        entry_fg_color = 'black'        
        button_bg_color = 'blue'        
        button_fg_color = 'white'       

        # Selecting the elements for Extract and Manipulation 
        self.partner_label = tk.Label(self.root, text='Partner')
        self.partner_street_label = tk.Label(self.root, text='Partner Street')
        self.partner_location_label = tk.Label(self.root, text='Partner ZIP_City_Country')
        self.invoice_num_label = tk.Label(self.root, text='Invoice Number')
        self.service_description_label = tk.Label(self.root, text='Service Description')
        self.service_amount_label = tk.Label(self.root, text='Amount')
        self.service_price_label = tk.Label(self.root, text='Single Price')
        self.payment_method_label = tk.Label(self.root, text='Payment Method')

        # Keys selection for Banks as Payment Methods
        self.payment_methods = {
            'Head Bank': {
                'Recipient': 'StreamBox GmbH',
                'Bank': 'Nord Europa Bank',
                'IBAN': 'HI23 451 243 2343',
                'BIC': 'AB23 3456'
            },
            'Quarter Bank': {
                'Recipient': 'StreamBox Services GmbH',
                'Bank': 'Central Moon Bank',
                'IBAN': 'GO65 451 243 2343',
                'BIC': 'KL65 3456'
            },
            'Regional Bank': {
                'Recipient': 'StreamBox Finance GmbH',
                'Bank': 'Suden Welt Bank',
                'IBAN': 'UN98 745 067 2343',
                'BIC': 'NG43 3456'
            }
        }

        # Input fields
        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_location_entry = tk.Entry(self.root)
        self.invoice_num_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_price_entry = tk.Entry(self.root)

        # Payment Method Dropdown Menu
        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set('Head Bank')
        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_method, *self.payment_methods.keys())
        
        self.create_button = tk.Button(self.root, text='Create Invoice', command=self.create_invoice)

        # Grid layout for the elements
        self.partner_label.grid(row=0, column=0, sticky='e', padx=10, pady=5)
        self.partner_entry.grid(row=0, column=1, padx=10, pady=5)
        self.partner_street_label.grid(row=1, column=0, sticky='e', padx=10, pady=5)
        self.partner_street_entry.grid(row=1, column=1, padx=10, pady=5)
        self.partner_location_label.grid(row=2, column=0, sticky='e', padx=10, pady=5)
        self.partner_location_entry.grid(row=2, column=1, padx=10, pady=5)
        self.invoice_num_label.grid(row=3, column=0, sticky='e', padx=10, pady=5)
        self.invoice_num_entry.grid(row=3, column=1, padx=10, pady=5)
        self.service_description_label.grid(row=4, column=0, sticky='e', padx=10, pady=5)
        self.service_description_entry.grid(row=4, column=1, padx=10, pady=5)
        self.service_amount_label.grid(row=5, column=0, sticky='e', padx=10, pady=5)
        self.service_amount_entry.grid(row=5, column=1, padx=10, pady=5)
        self.service_price_label.grid(row=6, column=0, sticky='e', padx=10, pady=5)
        self.service_price_entry.grid(row=6, column=1, padx=10, pady=5)
        self.payment_method_label.grid(row=7, column=0, sticky='e', padx=10, pady=5)
        self.payment_method_dropdown.grid(row=7, column=1, padx=10, pady=5)
        self.create_button.grid(row=8, column=0, columnspan=2, pady=10)

        # Apply color configurations
        self.partner_label.config(bg=label_bg_color, fg=label_fg_color)
        self.partner_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.partner_street_label.config(bg=label_bg_color, fg=label_fg_color)
        self.partner_street_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.partner_location_label.config(bg=label_bg_color, fg=label_fg_color)
        self.partner_location_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.invoice_num_label.config(bg=label_bg_color, fg=label_fg_color)
        self.invoice_num_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.service_description_label.config(bg=label_bg_color, fg=label_fg_color)
        self.service_description_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.service_amount_label.config(bg=label_bg_color, fg=label_fg_color)
        self.service_amount_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.service_price_label.config(bg=label_bg_color, fg=label_fg_color)
        self.service_price_entry.config(bg=entry_bg_color, fg=entry_fg_color)
        self.payment_method_label.config(bg=label_bg_color, fg=label_fg_color)
        self.payment_method_dropdown.config(bg=entry_bg_color, fg=entry_fg_color)
        self.create_button.config(bg=button_bg_color, fg=button_fg_color)

        self.root.mainloop()

    # Replace text in a paragraph
    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    def create_invoice(self):
        # Placeholder for invoice creation logic
        doc = docx.Document('template.docx')

        selected_payment_method = self.payment_methods[self.payment_method.get()]

        # Create a dictionary of text replacements
        try:
            replacements = {
                "[Date]": dt.datetime.today().strftime('%Y-%m-%d'),
                "[Partner]": self.partner_entry.get(),
                "[Partner Street]": self.partner_street_entry.get(),
                "[Partner ZIP_City_Country]": self.partner_location_entry.get(),
                "[Invoice Number]": self.invoice_num_entry.get(),
                "[Service Description]": self.service_description_entry.get(),
                "[Amount]": self.service_amount_entry.get(),
                "[Single Price]": f"${float(self.service_price_entry.get()):.2f}",
                "[Full Price]": f'${float(self.service_amount_entry.get()) * float(self.service_price_entry.get()):.2f}',
                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[IBAN]": selected_payment_method['IBAN'],
                "[BIC]": selected_payment_method['BIC'],
            }
        except ValueError:
            messagebox.showerror('Error', 'Invalid amount or price!')
            return
        
        # Replace text in paragraphs
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        # Replace text in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        # Save the filled document and convert to PDF
        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF documents', '*.pdf')])

        doc.save('filled.docx')
        # Assuming convert is a function to convert docx to pdf
        convert('filled.docx', save_path)

        messagebox.showinfo('Success', 'Invoice created and saved successfully')

if __name__ == '__main__':
    InvoiceMapper()
